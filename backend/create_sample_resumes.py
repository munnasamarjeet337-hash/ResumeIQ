import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_sample_docx(output_path="sample_resume_alex_morgan.docx"):
    doc = docx.Document()
    
    # Title / Name
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_after = Pt(2)
    name_run = title_p.add_run("Alex Morgan")
    name_run.font.name = "Arial"
    name_run.font.size = Pt(22)
    name_run.font.bold = True
    name_run.font.color.rgb = RGBColor(79, 70, 229) # Indigo

    # Contact line
    contact_p = doc.add_paragraph()
    contact_p.paragraph_format.space_after = Pt(12)
    contact_run = contact_p.add_run("Email: alex.morgan@example.com | Phone: (555) 234-5678 | San Francisco, CA\nLinkedIn: linkedin.com/in/alexmorgan | GitHub: github.com/alexmorgan")
    contact_run.font.name = "Arial"
    contact_run.font.size = Pt(9.5)
    contact_run.font.color.rgb = RGBColor(100, 116, 139)

    # Section 1: Professional Summary
    h1 = doc.add_paragraph()
    h1.paragraph_format.space_before = Pt(8)
    h1.paragraph_format.space_after = Pt(3)
    r1 = h1.add_run("PROFESSIONAL SUMMARY")
    r1.font.bold = True
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor(30, 41, 59)

    p_summary = doc.add_paragraph()
    p_summary.paragraph_format.space_after = Pt(10)
    p_summary.add_run(
        "Senior Full-Stack Engineer with 5+ years of experience architecting resilient cloud SaaS applications using React, Python, and Docker. "
        "Spearheaded microservices migration reducing API latency by 45% and scaled WebSocket streaming serving over 150k daily active users. "
        "Passionate about clean architecture, CI/CD automation, and mentoring high-performing engineering teams."
    )

    # Section 2: Technical Skills
    h2 = doc.add_paragraph()
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(3)
    r2 = h2.add_run("TECHNICAL SKILLS")
    r2.font.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(30, 41, 59)

    p_skills = doc.add_paragraph()
    p_skills.paragraph_format.space_after = Pt(10)
    p_skills.add_run("• Programming Languages: Python, JavaScript, TypeScript, SQL, HTML5, CSS3, Bash\n")
    p_skills.add_run("• Frameworks & Libraries: React, Node.js, Express, Flask, FastAPI, Tailwind CSS, Redux, Next.js\n")
    p_skills.add_run("• Cloud & DevOps: Docker, Kubernetes, AWS, PostgreSQL, MongoDB, Redis, Git, CI/CD, GitHub Actions\n")
    p_skills.add_run("• Competencies & Architecture: REST APIs, Microservices, System Design, Agile, Unit Testing (pytest, Jest)")

    # Section 3: Work Experience
    h3 = doc.add_paragraph()
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(3)
    r3 = h3.add_run("PROFESSIONAL EXPERIENCE")
    r3.font.bold = True
    r3.font.size = Pt(11)
    r3.font.color.rgb = RGBColor(30, 41, 59)

    p_exp1_head = doc.add_paragraph()
    p_exp1_head.paragraph_format.space_after = Pt(2)
    exp1_title = p_exp1_head.add_run("Senior Software Engineer | TechCorp Inc. (San Francisco, CA) — 2021 to Present\n")
    exp1_title.font.bold = True

    p_exp1 = doc.add_paragraph()
    p_exp1.paragraph_format.space_after = Pt(8)
    p_exp1.add_run("• Architected and developed high-throughput RESTful microservices in Python Flask and FastAPI backed by PostgreSQL and Redis.\n")
    p_exp1.add_run("• Scaled real-time WebSocket notification services to support 150,000+ daily active users with 99.99% uptime.\n")
    p_exp1.add_run("• Implemented automated CI/CD deployment pipelines using Docker and GitHub Actions, reducing release cycle time by 60%.\n")
    p_exp1.add_run("• Led sprint planning, conducted code reviews, and mentored 4 junior and mid-level full stack software engineers.")

    p_exp2_head = doc.add_paragraph()
    p_exp2_head.paragraph_format.space_after = Pt(2)
    exp2_title = p_exp2_head.add_run("Full-Stack Developer | CloudSystems LLC (Austin, TX) — 2019 to 2021\n")
    exp2_title.font.bold = True

    p_exp2 = doc.add_paragraph()
    p_exp2.paragraph_format.space_after = Pt(10)
    p_exp2.add_run("• Engineered responsive, WCAG-accessible web client interfaces using React, TypeScript, and Tailwind CSS.\n")
    p_exp2.add_run("• Designed optimized database schemas and relational queries in PostgreSQL, improving query response times by 40%.\n")
    p_exp2.add_run("• Built secure JWT authentication, session management, and RBAC authorization middleware.")

    # Section 4: Education
    h4 = doc.add_paragraph()
    h4.paragraph_format.space_before = Pt(8)
    h4.paragraph_format.space_after = Pt(3)
    r4 = h4.add_run("EDUCATION")
    r4.font.bold = True
    r4.font.size = Pt(11)
    r4.font.color.rgb = RGBColor(30, 41, 59)

    p_edu = doc.add_paragraph()
    p_edu.paragraph_format.space_after = Pt(8)
    p_edu.add_run("Bachelor of Science in Computer Science | University of California, Berkeley (2015 - 2019)\n")
    p_edu.add_run("Relevant Coursework: Data Structures & Algorithms, Distributed Systems, Database Management Systems.")

    doc.save(output_path)
    print(f"Sample DOCX resume created at: {output_path}")

if __name__ == '__main__':
    create_sample_docx()
