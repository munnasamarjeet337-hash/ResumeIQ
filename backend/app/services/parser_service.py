import os
import re
import io
import pypdf
import pdfplumber
import docx

class ResumeParserService:
    @staticmethod
    def extract_text_from_file(filepath):
        """Extracts cleaned raw text from PDF or DOCX file with graceful fallbacks."""
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
            
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext == '.pdf':
            return ResumeParserService._extract_from_pdf(filepath)
        elif ext in ('.docx', '.doc'):
            return ResumeParserService._extract_from_docx(filepath)
        else:
            # Fallback to plain text read
            with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()

    @staticmethod
    def _extract_from_pdf(filepath):
        text_content = []
        
        # 1. Primary: pdfplumber for high fidelity layout & text extraction
        try:
            with pdfplumber.open(filepath) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text(layout=True) or page.extract_text()
                    if page_text:
                        text_content.append(page_text)
            if text_content:
                return "\n\n".join(text_content)
        except Exception as e:
            print(f"[Parser] pdfplumber failed ({e}), falling back to pypdf...")

        # 2. Secondary fallback: pypdf
        try:
            reader = pypdf.PdfReader(filepath)
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted:
                    text_content.append(extracted)
            if text_content:
                return "\n\n".join(text_content)
        except Exception as e:
            print(f"[Parser] pypdf failed: {e}")

        raise ValueError("Could not extract readable text from PDF file. Ensure it is not password protected or purely scanned images.")

    @staticmethod
    def _extract_from_docx(filepath):
        try:
            doc = docx.Document(filepath)
            full_text = []
            
            # Paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())
                    
            # Tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        full_text.append(" | ".join(row_text))
                        
            return "\n\n".join(full_text)
        except Exception as e:
            raise ValueError(f"Could not parse DOCX file: {e}")

    @staticmethod
    def clean_text(raw_text):
        """Standardizes line breaks, whitespace, and special characters."""
        if not raw_text:
            return ""
        # Normalize non-breaking spaces and weird quotes
        text = raw_text.replace('\xa0', ' ').replace('\u2013', '-').replace('\u2014', '--')
        # Remove multiple consecutive blank lines
        text = re.sub(r'\n\s*\n', '\n\n', text)
        # Normalize multiple spaces
        text = re.sub(r'[ \t]+', ' ', text)
        return text.strip()

    @staticmethod
    def extract_contact_info(text):
        """Extracts email, phone number, LinkedIn, GitHub, and portfolio links."""
        info = {
            "email": None,
            "phone": None,
            "linkedin": None,
            "github": None,
            "portfolio": None,
            "candidate_name": None
        }
        
        # Email
        email_match = re.search(r'[a-zA-Z0-9_.+-]+@[a-zA-Z0-9-]+\.[a-zA-Z0-9-.]+', text)
        if email_match:
            info["email"] = email_match.group(0)
            
        # Phone
        phone_match = re.search(r'(?:(?:\+?1\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?', text)
        if phone_match:
            info["phone"] = phone_match.group(0).strip()
            
        # LinkedIn
        linkedin_match = re.search(r'(?:https?://)?(?:www\.)?linkedin\.com/(?:in|profile)/[a-zA-Z0-9_-]+', text, re.IGNORECASE)
        if linkedin_match:
            info["linkedin"] = linkedin_match.group(0)
            
        # GitHub
        github_match = re.search(r'(?:https?://)?(?:www\.)?github\.com/[a-zA-Z0-9_-]+', text, re.IGNORECASE)
        if github_match:
            info["github"] = github_match.group(0)

        # Candidate Name heuristic (first 2-3 capitalized words on the first non-empty lines)
        lines = [line.strip() for line in text.split('\n') if line.strip()]
        for line in lines[:5]:
            # Skip if line looks like contact info or header
            if '@' in line or 'http' in line or len(line) > 50:
                continue
            words = line.split()
            if 1 <= len(words) <= 4 and all(w[0].isupper() for w in words if w.isalpha()):
                info["candidate_name"] = line
                break

        return info

    @staticmethod
    def segment_sections(text):
        """Splits resume text into standardized semantic sections."""
        section_headers = {
            "summary": [r"summary", r"professional summary", r"executive summary", r"about me", r"profile", r"objective"],
            "experience": [r"experience", r"work experience", r"employment history", r"professional experience", r"work history", r"internships"],
            "education": [r"education", r"academic background", r"qualifications", r"degrees"],
            "skills": [r"skills", r"technical skills", r"core competencies", r"technologies", r"tools & technologies", r"key skills", r"skills & expertise"],
            "projects": [r"projects", r"personal projects", r"key projects", r"portfolio projects", r"academic projects"],
            "certifications": [r"certifications", r"certificates", r"licenses", r"credentials", r"courses"],
            "awards": [r"awards", r"honors", r"achievements", r"publications"]
        }
        
        sections = {key: "" for key in section_headers}
        sections["header"] = ""
        
        # Build regex for section header boundary detection
        header_patterns = []
        for sec_name, keywords in section_headers.items():
            for kw in keywords:
                header_patterns.append((sec_name, re.compile(rf'^(?:#+\s*)?(?:[0-9]+\.\s*)?{kw}\s*[:\-]?\s*$', re.IGNORECASE | re.MULTILINE)))

        lines = text.split('\n')
        current_section = "header"
        section_lines = {key: [] for key in section_headers}
        section_lines["header"] = []

        for line in lines:
            line_str = line.strip()
            if not line_str:
                continue

            matched_new_section = None
            for sec_name, pattern in header_patterns:
                if pattern.match(line_str):
                    matched_new_section = sec_name
                    break
            
            if matched_new_section:
                current_section = matched_new_section
            else:
                section_lines[current_section].append(line)

        for sec, lns in section_lines.items():
            sections[sec] = "\n".join(lns).strip()

        return sections
